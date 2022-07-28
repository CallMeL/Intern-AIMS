
import pandas as pd
import json
import docx
import os
def parser():
    name = os.listdir('./container/')[0]
    # print(name)
    name = './container/'+name
    wordDoc = docx.Document(name)
    if pd.DataFrame(wordDoc.tables).shape == (5,1):
        table0 = wordDoc.tables[0]
        data0 = [[cell.text for cell in row.cells] for row in table0.rows]
        df0 = pd.DataFrame(data0)
        df0.columns = ['key', 'value']
        df_table1 = df0
        df_table1 = df_table1.drop([0])

        table1 = wordDoc.tables[1]
        data1 = [[cell.text for cell in row.cells] for row in table1.rows]
        df1 = pd.DataFrame(data1)
        df1.columns = ['key', 'value']
        df_table2 = df1
        df_table2 = df_table2.drop([0])


        key = ['projectName','pubManager','pubDepositer','pubSupervisor',
                    'pubFinanceConsultant','absFinancier','pubGuarantor','absAssetName',
                    'pubIntegratedCost','pubFundsUse','pubRepaySource','pubProspectiveEarnings',
                    'pubDueTime','virtualAmount','pubPrincipalType','projectGrade','absProjectSource',
            'absIncomeRate','absPromanageRate','absTrustmanageRate','absCustodyFeeRate',
                    'absSuperviseRate','absFinanceRate','absInvestIncomeRate'
            ]
        df_out = pd.concat([df_table1,df_table2])
        if df_out.shape == (24,2):
            df_T = df_out.T
            df_T.columns = key
            df_T = df_T.drop('key')
            df_T = df_T.reset_index(drop=True)
            result = df_T.to_json(orient="records")
            return result
            parsed = json.loads(result)
            data = json.dumps(parsed, sort_keys=False, indent=4, ensure_ascii=False)
            print(data)
            return result
            return data
        else:
            return False
    else:
        return False